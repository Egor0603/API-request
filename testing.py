import requests as req
import docx
from docx.shared import Pt


arr_18reg = []
arr_of_types = []
arr_for_table = []

# запрос к API  и его обработка, как json
r = req.get("https://api.kontur.ru/dc.contacts/v1/cus")
r = r.json()

# Фильтр значений с 18 регионом
for i in r["cus"]:
    if i["region"] == "18":
        arr_18reg.append(i)
        if i["type"] not in arr_of_types:
            arr_of_types.append(i["type"])

# Создание списка списков с необходимыми для таблицы значениями
for i in arr_of_types:
    for j in range(len(arr_18reg)):
        inner_list = []
        if arr_18reg[j]["type"] == i:
            inner_list.append(arr_18reg[j]["type"])
            inner_list.append(arr_18reg[j]["code"])
            inner_list.append(arr_18reg[j]["name"])
            if arr_18reg[j]["soun"]:
                inner_list.append(arr_18reg[j]["soun"]["inn"])
            else:
                inner_list.append(" ")

            arr_for_table.append(inner_list)

# Запись в .docx  файл в виде таблицы
doc = docx.Document()

# Стилизация .docx файла
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)

# Словарь для хранения количества каждого типа
dictionary_of_types = {}
# Его заполнение
for i in arr_for_table:
    if i[0] not in dictionary_of_types:
        dictionary_of_types[i[0]] = 1
    elif i[0] in dictionary_of_types:
        dictionary_of_types[i[0]] += 1

# Вывод количества кадого типа в .dock файл
doc.add_paragraph(f"Общее - {len(arr_for_table)}")

for i in dictionary_of_types:
    doc.add_paragraph(f"{i} - {dictionary_of_types.get(i)}")

menu_table = doc.add_table(rows=1, cols=5)
menu_table.style = "Table Grid"

table_headers = ["№", "Тип", "Код", "Имя", "ИНН"]

# Заполнение заголовков таблицы
for i in range(len(table_headers)):
    menu_table.rows[0].cells[i].text = table_headers[i]

number = 1

# Заполнение таблицы
for type_of_org, code, name, inn in arr_for_table:
    row_cells = menu_table.add_row().cells
    row_cells[0].text = str(number)
    row_cells[1].text = type_of_org
    row_cells[2].text = str(code)
    row_cells[3].text = name
    row_cells[4].text = str(inn)
    number += 1

doc.save("table.docx")
